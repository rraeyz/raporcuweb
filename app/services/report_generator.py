import os
import markdown2
from datetime import datetime
from flask import current_app
import re
import traceback
import base64
from io import BytesIO

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError as e:
    WEASYPRINT_AVAILABLE = False
    print(f"WeasyPrint import hatası: {e}")

try:
    from latex2mathml.converter import convert as latex2mathml
    LATEX2MATHML_AVAILABLE = True
except ImportError:
    LATEX2MATHML_AVAILABLE = False

try:
    import matplotlib
    matplotlib.use('Agg')  # Headless mode
    import matplotlib.pyplot as plt
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

class ReportGenerator:
    """Rapor oluşturma ve dönüştürme servisi"""
    
    def __init__(self):
        self.upload_folder = current_app.config.get('UPLOAD_FOLDER', 'uploads')
        if not os.path.exists(self.upload_folder):
            os.makedirs(self.upload_folder)
    
    def latex_to_omml(self, latex_str):
        """LaTeX'i Office Math ML (OMML) formatına çevir - Word için düzenlenebilir"""
        try:
            from lxml import etree
            
            # Basit LaTeX ifadelerini OMML'e çevir
            # Üslü sayılar için: x^2 -> <m:sSup><m:e>x</m:e><m:sup>2</m:sup></m:sSup>
            
            # Office Math namespace
            ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
            
            # Basit üslü sayı parse (10^-5 gibi)
            if '^' in latex_str:
                parts = latex_str.split('^')
                if len(parts) == 2:
                    base = parts[0].strip()
                    exp = parts[1].strip('{}').strip()
                    
                    # OMML XML oluştur
                    omml = f'''<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
                        <m:sSup>
                            <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
                            <m:sup><m:r><m:t>{exp}</m:t></m:r></m:sup>
                        </m:sSup>
                    </m:oMath>'''
                    return omml
            
            # Basit metin olarak döndür
            return f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>{latex_str}</m:t></m:r></m:oMath>'
        except Exception as e:
            current_app.logger.error(f'OMML dönüştürme hatası: {e}')
            return None
    
    def process_latex_in_html(self, html, for_pdf=False):
        """HTML içindeki LaTeX formüllerini işle (markdown işleminden SONRA)"""
        try:
            # Önce display math ($$...$$) - daha uzun pattern önce
            def replace_display(match):
                latex = match.group(1).strip()
                
                if for_pdf and MATPLOTLIB_AVAILABLE:
                    try:
                        # Matplotlib ile LaTeX render (display) - orta boyut
                        fig, ax = plt.subplots(figsize=(5, 0.8))
                        ax.text(0.5, 0.5, f'${latex}$', fontsize=13, ha='center', va='center', transform=ax.transAxes)
                        ax.axis('off')
                        
                        buf = BytesIO()
                        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True, pad_inches=0.05)
                        plt.close(fig)
                        
                        img_data = base64.b64encode(buf.getvalue()).decode()
                        return f'<div style="text-align: center; margin: 12px 0;"><img src="data:image/png;base64,{img_data}" alt="Math: {latex}" style="max-width: 100%; height: auto;"/></div>'
                    except Exception as e:
                        current_app.logger.warning(f'PDF LaTeX render hatası: {e}')
                        return f'<div style="text-align: center;"><code>$${latex}$$</code></div>'
                
                elif not for_pdf and MATPLOTLIB_AVAILABLE:
                    # Word için Pandoc kullanılacak, buraya gelmemeli ama fallback için
                    try:
                        fig, ax = plt.subplots(figsize=(5, 0.8))
                        ax.text(0.5, 0.5, f'${latex}$', fontsize=13, ha='center', va='center', transform=ax.transAxes)
                        ax.axis('off')
                        
                        buf = BytesIO()
                        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', transparent=True, pad_inches=0.05)
                        plt.close(fig)
                        
                        img_data = base64.b64encode(buf.getvalue()).decode()
                        return f'<div style="text-align: center; margin: 12px 0;"><img src="data:image/png;base64,{img_data}" alt="Math: {latex}" style="max-width: 100%;"/></div>'
                    except Exception as e:
                        current_app.logger.warning(f'Word LaTeX render hatası: {e}')
                        return f'<div style="text-align: center;"><code>$${latex}$$</code></div>'
                
                return match.group(0)
            
            html = re.sub(r'\$\$(.+?)\$\$', replace_display, html, flags=re.DOTALL)
            
            # Sonra inline math ($...$) - daha kısa pattern
            def replace_inline(match):
                latex = match.group(1).strip()
                
                if for_pdf and MATPLOTLIB_AVAILABLE:
                    try:
                        # Matplotlib ile inline LaTeX render - metin boyutuna uygun (küçük)
                        fig, ax = plt.subplots(figsize=(3, 0.4))
                        ax.text(0.5, 0.5, f'${latex}$', fontsize=11, ha='center', va='center', transform=ax.transAxes)
                        ax.axis('off')
                        
                        buf = BytesIO()
                        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', transparent=True, pad_inches=0.02)
                        plt.close(fig)
                        
                        img_data = base64.b64encode(buf.getvalue()).decode()
                        return f'<img src="data:image/png;base64,{img_data}" style="vertical-align: middle; height: 1.1em; margin: 0 1px;" alt="{latex}"/>'
                    except Exception as e:
                        current_app.logger.warning(f'PDF LaTeX render hatası: {e}')
                        return f'<code>${latex}$</code>'
                
                elif not for_pdf and MATPLOTLIB_AVAILABLE:
                    # Word için Pandoc kullanılacak, buraya gelmemeli ama fallback için
                    try:
                        fig, ax = plt.subplots(figsize=(3, 0.4))
                        ax.text(0.5, 0.5, f'${latex}$', fontsize=11, ha='center', va='center', transform=ax.transAxes)
                        ax.axis('off')
                        
                        buf = BytesIO()
                        plt.savefig(buf, format='png', dpi=120, bbox_inches='tight', transparent=True, pad_inches=0.02)
                        plt.close(fig)
                        
                        img_data = base64.b64encode(buf.getvalue()).decode()
                        return f'<img src="data:image/png;base64,{img_data}" style="vertical-align: middle; height: 1.1em; margin: 0 1px;" alt="{latex}"/>'
                    except Exception as e:
                        current_app.logger.warning(f'Word LaTeX render hatası: {e}')
                        return f'<code>${latex}$</code>'
                
                return match.group(0)
            
            # Sadece metin içindeki $...$ işle (kod bloklarında değil)
            html = re.sub(r'(?<!<code>)\$([^\$\n]+?)\$(?!</code>)', replace_inline, html)
            
            return html
        except Exception as e:
            current_app.logger.error(f'LaTeX işleme hatası: {str(e)}')
            current_app.logger.error(traceback.format_exc())
            return html
    
    def convert_superscripts_subscripts(self, html):
        """Üslü ve altlı sayıları HTML sup/sub tag'lerine çevir"""
        try:
            # 10^-5 gibi formatları <sup> tag'ine çevir
            html = re.sub(r'(\d+)\^(-?\d+)', r'\1<sup>\2</sup>', html)
            
            # x^2, y^3 gibi formatları da çevir
            html = re.sub(r'([a-zA-Z])\ ?\^\ ?(\d+)', r'\1<sup>\2</sup>', html)
            
            # ×10^-5 gibi çarpma işaretli formatları da çevir  
            html = re.sub(r'(×|x)\s?10\^(-?\d+)', r'\1 10<sup>\2</sup>', html)
            
            # Alt simge için _kullanımı: H_2O -> H<sub>2</sub>O
            html = re.sub(r'([a-zA-Z])_(\d+)', r'\1<sub>\2</sub>', html)
            
            return html
        except Exception as e:
            current_app.logger.error(f'Superscript dönüşüm hatası: {str(e)}')
            return html
    
    def markdown_to_html(self, markdown_text, for_pdf=False):
        """Markdown'u HTML'e çevir - LaTeX desteği ile"""
        # Önce markdown'u HTML'e çevir
        html = markdown2.markdown(
            markdown_text,
            extras=['tables', 'fenced-code-blocks', 'break-on-newline', 'cuddled-lists', 'code-friendly', 'strike', 'task_list']
        )
        
        # Üslü ve altlı sayıları çevir
        html = self.convert_superscripts_subscripts(html)
        
        # Sonra HTML içindeki LaTeX'i işle
        html = self.process_latex_in_html(html, for_pdf=for_pdf)
        
        return html
    
    def generate_pdf(self, content, title, username):
        """Markdown içeriğinden PDF oluştur - Pandoc (LaTeX) ile öncelikli, WeasyPrint fallback"""
        try:
            # Önce Pandoc ile dene (en kaliteli LaTeX işleme)
            try:
                import subprocess
                
                # Dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'rapor_{username}_{timestamp}.pdf'
                filepath = os.path.join(self.upload_folder, filename)
                
                # Geçici markdown dosyası
                temp_md = os.path.join(self.upload_folder, f'temp_pdf_{timestamp}.md')
                with open(temp_md, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Pandoc komutu - xelatex engine ile PDF (Docker'da kurulu)
                cmd = [
                    'pandoc',
                    temp_md,
                    '-o', filepath,
                    '--from=markdown',
                    '--to=pdf',
                    '--pdf-engine=xelatex',
                    '-V', 'mainfont=Times New Roman',
                    '-V', 'fontsize=12pt',
                    '-V', 'geometry:margin=2.5cm',
                    '--standalone'
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                # Geçici dosyayı sil
                if os.path.exists(temp_md):
                    os.remove(temp_md)
                
                if result.returncode == 0 and os.path.exists(filepath):
                    file_size = os.path.getsize(filepath)
                    return filepath, file_size, None
                else:
                    current_app.logger.warning(f'Pandoc PDF hatası: {result.stderr}')
                    raise Exception("Pandoc PDF başarısız, WeasyPrint deneniyor")
                    
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
                current_app.logger.warning(f'Pandoc kullanılamıyor ({str(e)}), WeasyPrint ile devam ediliyor')
                
                # Pandoc başarısız, WeasyPrint ile devam et
                if not WEASYPRINT_AVAILABLE:
                    current_app.logger.error("WeasyPrint de kullanılamıyor!")
                    return None, None, "PDF oluşturulamadı: Pandoc ve WeasyPrint kullanılamadı."
                
                # Dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'rapor_{username}_{timestamp}.pdf'
                filepath = os.path.join(self.upload_folder, filename)
                
                # Markdown'u HTML'e çevir (PDF için LaTeX render)
                html_content = self.markdown_to_html(content, for_pdf=True)
            
            # Profesyonel PDF CSS stili - Times New Roman 12pt
            css_style = """
                @page {
                    size: A4;
                    margin: 2.5cm;
                }
                
                body {
                    font-family: 'Times New Roman', 'DejaVu Serif', serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    color: #000000;
                    text-align: justify;
                }
                
                h1 {
                    font-family: 'Times New Roman', 'DejaVu Serif', serif;
                    font-size: 18pt;
                    font-weight: bold;
                    color: #000000;
                    text-align: center;
                    margin-top: 0;
                    margin-bottom: 24px;
                    page-break-after: avoid;
                }
                
                h2 {
                    font-family: 'Times New Roman', 'DejaVu Serif', serif;
                    font-size: 14pt;
                    font-weight: bold;
                    color: #000000;
                    text-align: left;
                    margin-top: 18px;
                    margin-bottom: 10px;
                    page-break-after: avoid;
                }
                
                h3 {
                    font-family: 'Times New Roman', 'DejaVu Serif', serif;
                    font-size: 13pt;
                    font-weight: bold;
                    color: #000000;
                    text-align: left;
                    margin-top: 14px;
                    margin-bottom: 8px;
                    page-break-after: avoid;
                }
                
                h4, h5, h6 {
                    font-family: 'Times New Roman', 'DejaVu Serif', serif;
                    font-size: 12pt;
                    font-weight: bold;
                    color: #000000;
                    text-align: left;
                    margin-top: 12px;
                    margin-bottom: 6px;
                    page-break-after: avoid;
                }
                
                p {
                    margin-top: 0;
                    margin-bottom: 12px;
                    orphans: 3;
                    widows: 3;
                }
                
                /* Tablolar */
                table {
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                    page-break-inside: avoid;
                }
                
                th {
                    background-color: #3498db;
                    color: white;
                    font-weight: bold;
                    padding: 10px;
                    text-align: left;
                    border: 1px solid #2980b9;
                }
                
                td {
                    padding: 8px;
                    border: 1px solid #ddd;
                }
                
                tr:nth-child(even) {
                    background-color: #f9f9f9;
                }
                
                /* Listeler */
                ul, ol {
                    margin: 12px 0;
                    padding-left: 30px;
                }
                
                li {
                    margin-bottom: 8px;
                    line-height: 1.6;
                }
                
                /* Kod blokları */
                pre {
                    background-color: #f4f4f4;
                    border: 1px solid #ddd;
                    border-left: 4px solid #3498db;
                    padding: 12px;
                    margin: 15px 0;
                    overflow-x: auto;
                    page-break-inside: avoid;
                    font-family: 'Courier New', monospace;
                    font-size: 10pt;
                    line-height: 1.4;
                }
                
                code {
                    background-color: #f4f4f4;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                    font-size: 10pt;
                }
                
                /* Blockquote */
                blockquote {
                    border-left: 4px solid #3498db;
                    padding-left: 15px;
                    margin: 15px 0;
                    color: #555;
                    font-style: italic;
                }
                
                /* LaTeX denklemler */
                .math {
                    font-family: 'STIX Two Math', 'Latin Modern Math', serif;
                    font-size: 12pt;
                    text-align: center;
                    margin: 15px 0;
                }
                
                /* Bağlantılar */
                a {
                    color: #3498db;
                    text-decoration: none;
                }
                
                a:hover {
                    text-decoration: underline;
                }
                
                /* Tarih bilgisi */
                .report-meta {
                    color: #777;
                    font-size: 9pt;
                    text-align: center;
                    margin-bottom: 30px;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 15px;
                }
                
                /* Sayfa sonları */
                .page-break {
                    page-break-after: always;
                }
            """
            
            # HTML template - sadece içerik (başlık ve tarih HTML'den kaldırıldı)
            html_template = f"""
            <!DOCTYPE html>
            <html lang="tr">
            <head>
                <meta charset="UTF-8">
                <title>{title}</title>
                <style>{css_style}</style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # PDF oluştur - CSS objesi ile (WeasyPrint fallback)
            css_obj = CSS(string=css_style)
            html_obj = HTML(string=html_template)
            html_obj.write_pdf(target=filepath, stylesheets=[css_obj])
            
            # Dosya boyutunu al
            file_size = os.path.getsize(filepath)
            
            return filepath, file_size, None
            
        except Exception as e:
            current_app.logger.error(f'PDF oluşturma hatası: {str(e)}')
            current_app.logger.error(traceback.format_exc())
            return None, None, f'PDF oluşturma hatası: {str(e)}'
    
    def generate_word(self, content, title, username):
        """Markdown içeriğinden Word belgesi oluştur - Pandoc ile düzenlenebilir matematik"""
        try:
            # Önce Pandoc ile dene (LaTeX'i Word matematik formatına çevirir)
            try:
                import subprocess
                
                # Dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'rapor_{username}_{timestamp}.docx'
                filepath = os.path.join(self.upload_folder, filename)
                
                # Geçici markdown dosyası oluştur
                temp_md = os.path.join(self.upload_folder, f'temp_{timestamp}.md')
                with open(temp_md, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                # Pandoc ile Word'e dönüştür
                cmd = [
                    'pandoc',
                    temp_md,
                    '-o', filepath,
                    '--from=markdown',
                    '--to=docx',
                    '--standalone'
                ]
                
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                
                # Geçici dosyayı sil
                if os.path.exists(temp_md):
                    os.remove(temp_md)
                
                if result.returncode == 0 and os.path.exists(filepath):
                    file_size = os.path.getsize(filepath)
                    
                    # Word belgesini düzenle: Times New Roman, siyah başlıklar, hizalama
                    try:
                        from docx import Document
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.shared import Pt, RGBColor
                        
                        doc = Document(filepath)
                        
                        # Tüm paragrafları işle
                        first_h1_found = False
                        for para in doc.paragraphs:
                            # Times New Roman fontunu ayarla
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                            
                            # Başlıkları işle
                            if para.style.name.startswith('Heading'):
                                # İlk H1'i ortala
                                if para.style.name in ['Heading 1', 'Title', 'Heading1'] and not first_h1_found:
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    first_h1_found = True
                                else:
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                # Başlıkları siyah yap
                                for run in para.runs:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                                    run.font.name = 'Times New Roman'
                        
                        doc.save(filepath)
                    except Exception as e:
                        current_app.logger.warning(f'Word düzenleme hatası: {e}')
                        pass  # Düzenleme başarısız olsa da devam et
                    
                    return filepath, file_size, None
                else:
                    raise Exception(f"Pandoc hatası: {result.stderr}")
                    
            except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
                current_app.logger.warning(f'Pandoc kullanılamıyor ({str(e)}), htmldocx deneniyor')
                
                # Pandoc yoksa htmldocx ile dene
                try:
                    from htmldocx import HtmlToDocx
                    from docx import Document
                    
                    # Dosya adı oluştur
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f'rapor_{username}_{timestamp}.docx'
                    filepath = os.path.join(self.upload_folder, filename)
                    
                    # Markdown'u HTML'e çevir
                    html_content = self.markdown_to_html(content, for_pdf=False)
                    
                    # Word belgesi oluştur
                    doc = Document()
                    
                    # HTML'i Word'e dönüştür
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    new_parser = HtmlToDocx()
                    new_parser.add_html_to_document(html_content, doc)
                    
                    # İlk başlığı ortala
                    for para in doc.paragraphs[:3]:
                        if para.style.name in ['Heading 1', 'Title', 'Heading1']:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            break
                    
                    # Belgeyi kaydet
                    doc.save(filepath)
                    
                    file_size = os.path.getsize(filepath)
                    
                    if file_size < 5000:
                        current_app.logger.warning(f'htmldocx küçük dosya oluşturdu ({file_size} bytes), fallback kullanılacak')
                        raise Exception("htmldocx küçük dosya oluşturdu")
                    
                    return filepath, file_size, None
                    
                except (ImportError, Exception) as e:
                    current_app.logger.warning(f'htmldocx başarısız ({str(e)}), fallback yöntem kullanılıyor')
                # Fallback: geliştirilmiş markdown parsing
                from docx import Document
                from docx.shared import Pt, RGBColor, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f'rapor_{username}_{timestamp}.docx'
                filepath = os.path.join(self.upload_folder, filename)
                
                # Word belgesi oluştur
                doc = Document()
                
                # Markdown içeriğini satır satır işle
                first_h1 = True  # İlk H1 başlığını takip et
                lines = content.split('\n')
                i = 0
                while i < len(lines):
                    line = lines[i].strip()
                    
                    if not line:
                        i += 1
                        continue
                    
                    # Başlıklar
                    if line.startswith('# '):
                        heading = doc.add_heading(line[2:], level=1)
                        if first_h1:
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            first_h1 = False
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('#### '):
                        doc.add_heading(line[5:], level=4)
                    # Liste öğeleri
                    elif line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    elif line.startswith(tuple(f'{j}. ' for j in range(1, 100))):
                        doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
                    # Normal paragraf
                    else:
                        # Kalın metin (**text**) işle
                        paragraph = doc.add_paragraph()
                        parts = line.split('**')
                        for idx, part in enumerate(parts):
                            run = paragraph.add_run(part)
                            if idx % 2 == 1:  # Kalın yapılacak kısımlar
                                run.bold = True
                    
                    i += 1
                
                # Belgeyi kaydet
                doc.save(filepath)
                
                # Dosya boyutunu al
                file_size = os.path.getsize(filepath)
                
                return filepath, file_size, None
            
        except Exception as e:
            current_app.logger.error(f'Word oluşturma hatası: {str(e)}')
            current_app.logger.error(traceback.format_exc())
            return None, None, f'Word oluşturma hatası: {str(e)}'
    
    def get_file_path(self, filename):
        """Dosya yolunu döndür"""
        return os.path.join(self.upload_folder, filename)
    
    def delete_file(self, filepath):
        """Dosyayı sil"""
        try:
            if os.path.exists(filepath):
                os.remove(filepath)
                return True
            return False
        except Exception as e:
            current_app.logger.error(f'Dosya silme hatası: {str(e)}')
            return False
